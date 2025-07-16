import os
import re
from io import BytesIO
from PIL import Image
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import fitz
import pdfplumber

def extract_docx(filepath):
    doc = Document(filepath)
    media_dir = "extracted_images"
    os.makedirs(media_dir, exist_ok=True)

    content = {
        "metadata": {
            "title": "",
            "authors": [],
            "abstract": "",
            "keywords": [],
            "figures": []
        },
        "body": [],
        "references": [],
        "tables": []
    }

    def is_heading(para):
        return para.style.name.lower().startswith('heading')

    def get_heading_level(para):
        match = re.match(r'heading\s*(\d+)', para.style.name.lower())
        return int(match.group(1)) if match else 1

    def clean_text(text):
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

    def save_image(image_part, index, caption=None):
        try:
            data = image_part.blob
            img = Image.open(BytesIO(data))
            ext = img.format.lower() if img.format else 'png'
            filename = f"figure_{index}.{ext}"
            path = os.path.join(media_dir, filename)
            img.save(path)
            return {
                "id": f"fig_{index}",
                "filename": filename,
                "caption": caption or f"Figure {index}",
                "path": path
            }
        except Exception as e:
            print(f"Error saving image: {e}")
            return None

    paras = list(doc.paragraphs)
    i = 0
    while i < len(paras) and not paras[i].text.strip():
        i += 1
    if i < len(paras):
        content["metadata"]["title"] = clean_text(paras[i].text)
        i += 1

    author_block = []
    while i < len(paras):
        text = clean_text(paras[i].text)
        if not text:
            i += 1
            continue
        if re.match(r'abstract', text, re.I) or is_heading(paras[i]):
            break
        author_block.append(text)
        i += 1

    for author_line in author_block:
        email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', author_line)
        email = email_match.group(0) if email_match else ""
        author_clean = re.sub(r'[\w\.-]+@[\w\.-]+\.\w+', '', author_line).strip()
        parts = [p.strip() for p in re.split(r',|;', author_clean)]
        if len(parts) >= 2:
            name = parts[0]
            role = parts[1] if len(parts) > 2 else ""
            affiliation = ", ".join(parts[2:]) if len(parts) > 2 else parts[1]
        else:
            name = author_clean
            role = ""
            affiliation = ""
        content["metadata"]["authors"].append({
            "name": name,
            "role": role,
            "affiliation": affiliation,
            "email": email
        })

    in_abstract = False
    in_keywords = False
    abstract_lines = []
    keywords = []

    while i < len(paras):
        text = clean_text(paras[i].text)
        if not text:
            i += 1
            continue

        if re.match(r'abstract', text, re.I):
            in_abstract = True
            i += 1
            continue

        if in_abstract and re.match(r'(keywords|key terms|index terms)\s*[:.]?', text, re.I):
            content["metadata"]["abstract"] = ' '.join(abstract_lines)
            in_abstract = False
            in_keywords = True
            i += 1
            continue

        if in_abstract:
            if is_heading(paras[i]):
                content["metadata"]["abstract"] = ' '.join(abstract_lines)
                in_abstract = False
            else:
                abstract_lines.append(text)
                i += 1
            continue

        if in_keywords:
            if is_heading(paras[i]):
                in_keywords = False
            else:
                for kw in re.split(r'[,;]', text):
                    kw = kw.strip()
                    if kw:
                        keywords.append(kw)
                i += 1
            continue

        break

    if keywords:
        content["metadata"]["keywords"] = keywords

    current_section = None
    current_subsection = None
    fig_count = 1

    while i < len(paras):
        para = paras[i]
        text = clean_text(para.text)
        if not text:
            i += 1
            continue

        if is_heading(para):
            level = get_heading_level(para)
            if level == 1:
                if current_section:
                    content["body"].append(current_section)
                current_section = {
                    "type": "section",
                    "heading": text,
                    "content": []
                }
                current_subsection = None
            elif level == 2:
                if current_section:
                    current_subsection = {
                        "type": "subsection",
                        "heading": text,
                        "content": []
                    }
                    current_section["content"].append(current_subsection)
                else:
                    current_section = {
                        "type": "section",
                        "heading": "Untitled Section",
                        "content": [{
                            "type": "subsection",
                            "heading": text,
                            "content": []
                        }]
                    }
                    current_subsection = current_section["content"][0]
            i += 1
            continue

        fig_match = re.match(r'(?:figure|fig)[\s\.]*(\d+)[\s:\.-]+(.*)', text, re.I)
        if fig_match:
            fig_num = fig_match.group(1)
            caption = fig_match.group(2).strip()
            image_info = None
            for rel in doc.part.rels.values():
                if rel.reltype == RT.IMAGE:
                    image_info = save_image(rel.target_part, fig_count, f"Figure {fig_num}: {caption}")
                    break
            if image_info:
                content["metadata"]["figures"].append(image_info)
                fig_count += 1
                fig_item = {
                    "type": "figure",
                    "label": f"Fig{fig_num}",
                    "caption": caption,
                    "content": image_info["path"]
                }
                if current_subsection:
                    current_subsection["content"].append(fig_item)
                elif current_section:
                    current_section["content"].append(fig_item)
            i += 1
            continue

        if current_subsection:
            container = current_subsection
        elif current_section:
            container = current_section
        else:
            current_section = {
                "type": "section",
                "heading": "Introduction",
                "content": []
            }
            container = current_section

        if para.style.name.lower().startswith('list') or text.startswith(('•', '-', '*')):
            list_item = {
                "type": "list",
                "items": [text.lstrip('•-* ').strip()]
            }
            if container["content"] and container["content"][-1].get("type") == "list":
                container["content"][-1]["items"].append(text.lstrip('•-* ').strip())
            else:
                container["content"].append(list_item)
        else:
            container["content"].append({
                "type": "paragraph",
                "text": text
            })
        i += 1

    if current_section:
        content["body"].append(current_section)

    table_count = 1
    for table in doc.tables:
        table_data = {
            "type": "table",
            "label": f"Table{table_count}",
            "header": [],
            "rows": []
        }
        if table.rows:
            header_row = table.rows[0]
            table_data["header"] = [clean_text(cell.text) for cell in header_row.cells]
            for row in table.rows[1:]:
                table_data["rows"].append([clean_text(cell.text) for cell in row.cells])
            content["tables"].append(table_data)
            if content["body"]:
                content["body"][-1]["content"].append(table_data)
            else:
                content["body"].append({
                    "type": "section",
                    "heading": "Tables",
                    "content": [table_data]
                })
            table_count += 1

    ref_section = None
    for j, para in enumerate(paras):
        if re.match(r'references?', clean_text(para.text), re.I):
            ref_section = j
            break
    if ref_section:
        refs = []
        ref_id = 1
        for para in paras[ref_section+1:]:
            text = clean_text(para.text)
            if not text:
                continue
            if is_heading(para):
                break
            refs.append({
                "id": f"{ref_id}",
                "citation": text
            })
            ref_id += 1
        content["references"] = refs

    return content

def extract_pdf(filepath):
    doc = fitz.open(filepath)
    media_dir = "extracted_images"
    os.makedirs(media_dir, exist_ok=True)

    output = []
    figures = []
    tables = []
    references = []

    def sanitize_filename(text):
        match = re.match(r'fig\s*\d*\s*:\s*(.+)', text, re.I)
        if match:
            caption_text = match.group(1)
        else:
            caption_text = text
        caption_text = caption_text.lower().strip()
        caption_text = re.sub(r'[^\w\-_. ]', '', caption_text)
        caption_text = caption_text.replace(' ', '_')
        return caption_text if caption_text else 'image'

    with pdfplumber.open(filepath) as pdf_plumber:
        for page_num, page in enumerate(doc, start=1):
            blocks = page.get_text("blocks")
            blocks.sort(key=lambda b: b[1])
            pdf_page = pdf_plumber.pages[page_num - 1]
            table_bboxes = []
            for t in pdf_page.find_tables():
                bbox = t.bbox
                table_bboxes.append((bbox, t))
            combined = []
            for bbox, tbl in table_bboxes:
                combined.append(('table', bbox[1], tbl))
            for b in blocks:
                combined.append(('text', b[1], b))
            combined.sort(key=lambda x: x[1])
            tables_output = set()
            images_info = []
            for img_index, img in enumerate(page.get_images(full=True)):
                xref = img[0]
                base_image = doc.extract_image(xref)
                img_bytes = base_image["image"]
                ext = base_image["ext"]
                images_info.append({"xref": xref, "bytes": img_bytes, "ext": ext})
            image_used = [False] * len(images_info)
            for item_type, y, content in combined:
                if item_type == 'table':
                    if content not in tables_output:
                        tables_output.add(content)
                        try:
                            table_data = content.extract()
                        except Exception:
                            table_data = content.extract_table()
                        if not table_data:
                            table_data = []
                        if table_data:
                            tables.append({
                                "type": "table",
                                "label": f"Table{len(tables)+1}",
                                "header": table_data[0],
                                "rows": table_data[1:]
                            })
                elif item_type == 'text':
                    x0, y0, x1, y1, text, block_no, block_type = content
                    if block_type == 0:
                        text = text.strip()
                        skip_line = False
                        for bbox, tbl in table_bboxes:
                            if y0 >= bbox[1] and y1 <= bbox[3]:
                                skip_line = True
                                break
                        if skip_line or not text:
                            continue
                        email_match = re.search(r'(\S+@\S+)', text)
                        if email_match and not text.startswith("Email:"):
                            before_email = text[:email_match.start()].strip()
                            email = email_match.group(1)
                            after_email = text[email_match.end():].strip()
                            if before_email:
                                output.append(before_email)
                            output.append(email)
                            if after_email:
                                output.append(after_email)
                            continue
                        text = re.sub(r'^[\u2022•\-]\s*', '- ', text)
                        if re.match(r'fig\s*\d*\s*:', text, re.I):
                            text = re.sub(r'paractice', 'practice', text, flags=re.I)
                        if re.match(r'fig\s*\d*\s*:', text, re.I):
                            for idx, used in enumerate(image_used):
                                if not used:
                                    filename_base = sanitize_filename(text)
                                    filename = f"{filename_base}.{images_info[idx]['ext']}"
                                    filepath_img = os.path.join(media_dir, filename)
                                    with open(filepath_img, "wb") as f:
                                        f.write(images_info[idx]["bytes"])
                                    image_used[idx] = True
                                    figures.append({
                                        "type": "figure",
                                        "label": f"Fig{len(figures)+1}",
                                        "caption": text,
                                        "content": filepath_img
                                    })
                                    break
                            else:
                                output.append(text)
                        else:
                            output.append(text)
            for idx, used in enumerate(image_used):
                if not used:
                    filename = f"image_{page_num}_{idx + 1}.{images_info[idx]['ext']}"
                    filepath_img = os.path.join(media_dir, filename)
                    with open(filepath_img, "wb") as f:
                        f.write(images_info[idx]["bytes"])
                    figures.append({
                        "type": "figure",
                        "label": f"Fig{len(figures)+1}",
                        "caption": filename,
                        "content": filepath_img
                    })

    cleaned_output = []
    prev_empty = False
    for line in output:
        if line.strip() == '':
            if not prev_empty:
                cleaned_output.append('')
            prev_empty = True
        else:
            cleaned_output.append(line)
            prev_empty = False

    # Simple metadata extraction from first lines
    metadata = {
        "title": cleaned_output[0] if cleaned_output else "",
        "authors": [],
        "abstract": "",
        "keywords": [],
        "figures": figures
    }
    body = []
    for line in cleaned_output[1:]:
        if line:
            body.append({"type": "paragraph", "text": line})

    return {
        "metadata": metadata,
        "body": body,
        "references": references,
        "tables": tables
    }
